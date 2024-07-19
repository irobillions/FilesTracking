import os
from openpyxl import load_workbook
from docx import Document
from tracking.summarize import *
import aspose.words as aw
import fitz  # PyMuPDF


class FileInfo:
    """
    Classe pour stocker les informations sur un fichier.
    """
    def __init__(self, name, path, summary):
        self.name = name
        self.path = path
        self.summary = summary


def file_contains_keyword_pdf(file_path, keyword):
    """
    Vérifie si un fichier PDF contient un mot clé.

    Args:
        file_path (str): Chemin du fichier PDF.
        keyword (str): Mot clé à rechercher.

    Returns:
        bool: True si le mot clé est trouvé, sinon False.
    """
    try:
        if file_path.startswith('~$') or file_path.startswith('._') or '~$' in file_path or '._ ' in file_path:
            return False

        if not file_path.lower().endswith(".pdf") or not os.path.isfile(file_path):
            return False
        pdf_document = fitz.open(file_path)
        for page_num in range(pdf_document.page_count):
            page = pdf_document.load_page(page_num)
            text = page.get_text("text")
            if keyword in text:
                return True
    except Exception as e:
        print(f"Erreur lors de la lecture du fichier PDF: {file_path}, Erreur: {e}")
    return False


def file_contains_keyword_txt(file_path, keyword):
    """
    Vérifie si un fichier texte contient un mot clé.

    Args:
        file_path (str): Chemin du fichier texte.
        keyword (str): Mot clé à rechercher.

    Returns:
        bool: True si le mot clé est trouvé, sinon False.
    """
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
            for line in file:
                if keyword in line:
                    return True
    except Exception as e:
        print(f"Erreur lors de la lecture du fichier texte: {file_path}, Erreur: {e}")
    return False


def file_contains_keyword_docx(file_path, keyword):
    """
    Vérifie si un fichier DOCX contient un mot clé.

    Args:
        file_path (str): Chemin du fichier DOCX.
        keyword (str): Mot clé à rechercher.

    Returns:
        bool: True si le mot clé est trouvé, sinon False.
    """
    try:
        if file_path.startswith('~$') or '~$' in file_path or '._ ' in file_path:
            return False
        doc = Document(file_path)
        if len(doc.paragraphs) <= 0:
            return False

        for para in doc.paragraphs:
            if keyword in para.text:
                return True
    except Exception as e:
        print(f"Erreur lors de la lecture du fichier DOCX: {file_path}, Erreur: {e}")
    return False


def file_contains_keyword_xlsx(file_path, keyword):
    """
    Vérifie si un fichier XLSX contient un mot clé.

    Args:
        file_path (str): Chemin du fichier XLSX.
        keyword (str): Mot clé à rechercher.

    Returns:
        bool: True si le mot clé est trouvé, sinon False.
    """
    try:
        workbook = load_workbook(file_path, read_only=True)
        for sheet in workbook:
            for row in sheet.iter_rows(values_only=True):
                for cell in row:
                    if cell and keyword in str(cell):
                        return True
    except Exception as e:
        print(f"Erreur lors de la lecture du fichier XLSX: {file_path}, Erreur: {e}")
    return False


def file_contains_keyword_doc(file_path, keyword):
    """
    Vérifie si un fichier DOC contient un mot clé.

    Args:
        file_path (str): Chemin du fichier DOC.
        keyword (str): Mot clé à rechercher.

    Returns:
        bool: True si le mot clé est trouvé, sinon False.
    """
    try:
        if  '~$' in file_path or '._ ' in file_path:
            return False
        doc = aw.Document(file_path)
        for run in doc.get_child_nodes(aw.NodeType.RUN, True):
            if keyword in run.get_text():
                return True
        else:
            return False
    except Exception as e:
        print(f"Erreur lors de la lecture du fichier DOC: {file_path}, Erreur: {e}")
    return False


def search_files_contain_keyword(directory, keyword):
    """
    Cherche des fichiers dans un répertoire qui contiennent un mot clé dans leur contenu.

    Args:
        directory (str): Chemin du répertoire à parcourir.
        keyword (str): Mot clé à rechercher dans les fichiers.

    Returns:
        list: Liste des objets FileInfo contenant les fichiers trouvés.
    """
    results = []
    for root, _, files in os.walk(directory):
        for file in files:
            file_path = os.path.join(root, file)
            extension = file.split('.')[-1].lower()
            contains_keyword = False
            if extension == 'txt':
                contains_keyword = file_contains_keyword_txt(file_path, keyword)
            elif extension == 'docx':
                contains_keyword = file_contains_keyword_docx(file_path, keyword)
            elif extension == 'xlsx':
                contains_keyword = file_contains_keyword_xlsx(file_path, keyword)
            elif extension == 'doc':
                contains_keyword = file_contains_keyword_doc(file_path, keyword)
            elif extension == 'pdf':
                contains_keyword = file_contains_keyword_pdf(file_path, keyword)
            if contains_keyword:
                summary = summarize_file(file_path)
                results.append(FileInfo(file, file_path, summary))
    return results


def search_files_with_keyword(directory, keyword):
    """
    Cherche des fichiers dans un répertoire qui contiennent un mot clé dans leur nom.

    Args:
        directory (str): Chemin du répertoire à parcourir.
        keyword (str): Mot clé à rechercher dans les noms de fichiers.

    Returns:
        list: Liste des objets FileInfo contenant les fichiers trouvés.
    """
    results = []
    for root, _, files in os.walk(directory):
        for file in files:
            if keyword in file:
                full_path = os.path.join(root, file)
                summary = summarize_file(full_path)
                results.append(FileInfo(file, full_path, summary))
    return results