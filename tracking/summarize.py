import os
import csv
from openpyxl import load_workbook
from docx import Document as DocxDocument
import fitz  # PyMuPDF
import aspose.words as aw
from sumy.parsers.plaintext import PlaintextParser
from sumy.nlp.tokenizers import Tokenizer
from sumy.summarizers.lsa import LsaSummarizer
from langdetect import detect
import nltk

nltk.download('punkt')


def summarize_text(text, language='english'):
    """
    Génère un résumé intelligent du texte donné.
    """
    try:
        parser = PlaintextParser.from_string(text, Tokenizer(language))
        summarizer = LsaSummarizer()
        summary = summarizer(parser.document, 2)  # Résumé avec 2 phrases
        return ' '.join(str(sentence) for sentence in summary)
    except Exception as e:
        return f"Erreur lors du résumé: {e}"


def summarize_txt(file_path):
    """
    Résume le contenu d'un fichier texte en extrayant les premières lignes.
    """
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
            text = file.read(2000)  # Lire les premiers 2000 caractères

        if len(text) <= 0:
            return "fichier vide"
        language = detect(text)
        summary = summarize_text(text, language)
        return f"Type de fichier: Texte\nContenu: {summary}"
    except Exception as e:
        return f"Erreur lors de la lecture du fichier texte: {e}"


def summarize_pdf(file_path):
    """
    Résume le contenu d'un fichier PDF en extrayant le texte des premières pages.
    """
    try:
        if '~$' in file_path or '._ ' in file_path:
            return "Fichier temporaire ignoré"

        if not file_path.lower().endswith(".pdf") or not os.path.isfile(file_path):
            return "Fichier temporaire ignoré"

        pdf_document = fitz.open(file_path)
        text = ''
        for page_num in range(min(3, pdf_document.page_count)):
            page = pdf_document.load_page(page_num)
            text += page.get_text("text")
            if len(text) > 2000:
                break
        language = detect(text)
        summary = summarize_text(text[:2000], language)
        return f"Type de fichier: PDF\nContenu: {summary}"
    except Exception as e:
        return f"Erreur lors de la lecture du fichier PDF: {e}"


def summarize_docx(file_path):
    """
    Résume le contenu d'un fichier DOCX en extrayant les premiers paragraphes.
    """
    try:
        if file_path.startswith('~$') or '~$' in file_path or '._ ' in file_path:
            return "Fichier temporaire ignoré."
        if file_path.startswith('~$'):
            return "Fichier temporaire ignoré."
        doc = DocxDocument(file_path)
        text = ' '.join([para.text for para in doc.paragraphs[:20]])
        if len(text) <= 0:
            return "fichier vide"

        language = detect(text)
        summary = summarize_text(text, language)
        return f"Type de fichier: DOCX\nContenu: {summary}"
    except Exception as e:
        return f"Erreur lors de la lecture du fichier DOCX: {e}"


def summarize_xlsx(file_path):
    """
    Résume le contenu d'un fichier XLSX en extrayant les premières cellules des premières feuilles.
    """
    try:
        workbook = load_workbook(file_path, read_only=True)
        sheet = workbook.active
        cells = []
        for row in sheet.iter_rows(min_row=1, max_row=10, max_col=5, values_only=True):
            cells.append(' '.join([str(cell) for cell in row if cell is not None]))
        text = ' '.join(cells)
        language = detect(text)
        summary = summarize_text(text, language)
        return f"Type de fichier: XLSX\nContenu: {summary}"
    except Exception as e:
        return f"Erreur lors de la lecture du fichier XLSX: {e}"


def summarize_doc(file_path):
    """
    Résume le contenu d'un fichier DOC en extrayant les premiers paragraphes.
    """
    try:
        doc = aw.Document(file_path)
        paragraphs = []
        for i, run in enumerate(doc.get_child_nodes(aw.NodeType.RUN, True)):
            if i >= 20:
                break
            paragraphs.append(run.get_text())
        text = ' '.join(paragraphs)

        if len(text) <= 0:
            return "Fichier Vide"

        language = detect(text)
        summary = summarize_text(text, language)
        return f"Type de fichier: DOC\nContenu: {summary}"
    except Exception as e:
        return f"Erreur lors de la lecture du fichier DOC: {e}"


def summarize_file(file_path):
    """
    Génère un résumé du contenu du fichier en fonction de son type.
    """
    extension = file_path.split('.')[-1].lower()
    if extension == 'txt':
        return summarize_txt(file_path)
    elif extension == 'pdf':
        return summarize_pdf(file_path)
    elif extension == 'docx':
        return summarize_docx(file_path)
    elif extension == 'xlsx':
        return summarize_xlsx(file_path)
    elif extension == 'doc':
        return summarize_doc(file_path)
    else:
        return "Type de fichier non supporté."
